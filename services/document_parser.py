# services/document_parser.py - 文档解析服务
# 从 PDF / DOCX 中提取文本 + 图片，并智能拆分为知识库 Q&A 条目
import os
import re
import logging

logger = logging.getLogger(__name__)


# ============================================================
# 文档文本 + 图片提取
# ============================================================

def _get_image_dir():
    """返回知识库图片存储目录（不存在则创建）"""
    from config import KB_IMAGE_DIR
    os.makedirs(KB_IMAGE_DIR, exist_ok=True)
    return KB_IMAGE_DIR


def _infer_image_ext(blip_element):
    """从 blip 元素的 extension 属性推断图片扩展名"""
    ext = blip_element.get("extension", "")
    ext_map = {
        "jpeg": "jpg", "png": "png", "gif": "gif",
        "bmp": "bmp", "tiff": "tiff", "emf": "emf",
        "wmf": "wmf", "pict": "pict",
    }
    return ext_map.get(ext.lower(), "png")


def extract_images_from_docx(file_path: str) -> list[tuple[bytes, str]]:
    """
    从 Word 文档提取所有嵌入图片。
    返回 [(image_bytes, filename), ...]
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    image_dir = _get_image_dir()

    # 收集文档中所有图片部分（通过 rels）
    seen = set()
    raw_images = []  # (bytes, original_ext)

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_part = rel.target_part
            content_type = image_part.content_type
            if "image" not in content_type:
                continue
            ext = content_type.split("/")[-1]
            if ext == "jpeg":
                ext = "jpg"
            raw_images.append((image_part.blob, ext))

    # 去重：按内容哈希避免重复保存
    image_hashes = set()
    results = []
    for idx, (img_bytes, ext) in enumerate(raw_images, 1):
        h = hash(img_bytes)
        if h in image_hashes:
            continue
        image_hashes.add(h)

        filename = f"kb_docx_{idx}.{ext}"
        filepath = os.path.join(image_dir, filename)
        with open(filepath, "wb") as f:
            f.write(img_bytes)
        results.append((img_bytes, filename))
        logger.info(f"DOCX 图片提取: {filename} ({len(img_bytes)} bytes)")

    return results


def _extract_images_from_pdf_pymupdf(file_path: str) -> list[tuple[bytes, str]]:
    """使用 PyMuPDF (fitz) 从 PDF 提取图片"""
    import fitz

    doc = fitz.open(file_path)
    image_dir = _get_image_dir()

    seen = set()
    raw_images = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images(full=True)
        for img_info in image_list:
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                ext = base_image["ext"]
                raw_images.append((image_bytes, ext, page_num + 1))
            except Exception as e:
                logger.warning(f"PDF 图片提取失败 xref={xref}: {e}")
                continue

    doc.close()

    image_hashes = set()
    results = []
    for idx, (img_bytes, ext, page_num) in enumerate(raw_images, 1):
        h = hash(img_bytes)
        if h in image_hashes:
            continue
        image_hashes.add(h)

        filename = f"kb_pdf_p{page_num}_{idx}.{ext}"
        filepath = os.path.join(image_dir, filename)
        with open(filepath, "wb") as f:
            f.write(img_bytes)
        results.append((img_bytes, filename))
        logger.info(f"PDF 图片提取: {filename} ({len(img_bytes)} bytes)")

    return results


def _extract_images_from_pdf_pdfplumber(file_path: str) -> list[tuple[bytes, str]]:
    """使用 pdfplumber 从 PDF 提取图片（备选方案）"""
    import pdfplumber

    doc = pdfplumber.open(file_path)
    image_dir = _get_image_dir()

    seen = set()
    raw_images = []

    for page_num, page in enumerate(doc.pages, 1):
        images = page.images
        for img_info in images:
            img_name = img_info.get("name", f"img_{len(raw_images) + 1}")
            # pdfplumber 不直接提供 image bytes，需要其他方式
            # 跳过此方案，仅作为占位
            pass

    doc.close()
    return []


def extract_images_from_pdf(file_path: str) -> list[tuple[bytes, str]]:
    """
    从 PDF 提取所有图片。
    优先使用 PyMuPDF (fitz)，备选 pdfplumber。
    """
    # 方案 1: PyMuPDF
    try:
        import fitz
        return _extract_images_from_pdf_pymupdf(file_path)
    except ImportError:
        pass

    # 方案 2: pdfplumber（当前仅占位）
    try:
        import pdfplumber
        return _extract_images_from_pdf_pdfplumber(file_path)
    except ImportError:
        logger.warning("PDF 图片提取：未安装 PyMuPDF 或 pdfplumber，图片将被跳过")

    return []


def extract_text_from_docx(file_path: str) -> tuple[str, list[tuple[bytes, str]]]:
    """
    从 Word 文档提取纯文本（保留段落结构）和图片。
    返回 (text, [(image_bytes, filename), ...])
    """
    from docx import Document

    doc = Document(file_path)
    paragraphs = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    text = "\n".join(paragraphs)
    images = extract_images_from_docx(file_path)
    return text, images


def extract_text_from_pdf(file_path: str) -> tuple[str, list[tuple[bytes, str]]]:
    """
    从 PDF 文档提取纯文本和图片。
    返回 (text, [(image_bytes, filename), ...])
    """
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages_text.append(text)

    text = "\n\n".join(pages_text)
    images = extract_images_from_pdf(file_path)
    return text, images


# ============================================================
# 智能 Q&A 拆分
# ============================================================

# 匹配 Q:/A: 或 问题：/答案：等显式问答模式
_QA_PATTERNS = [
    (re.compile(
        r"(?:Q\s*[:：]\s*|问题\s*[:：]\s*|问\s*[:：]\s*)(.*?)"
        r"(?:A\s*[:：]\s*|答案\s*[:：]\s*|答\s*[:：]\s*)(.*?)"
        r"(?=(?:Q\s*[:：]\s*|问题\s*[:：]\s*|问\s*[:：]\s*|$))",
        re.DOTALL | re.IGNORECASE
    ), "qa"),
    (re.compile(
        r"(?:Question\s*[:：]\s*|Q\.?\s*[:：]?\s*)(.*?)"
        r"(?:Answer\s*[:：]\s*|A\.?\s*[:：]?\s*)(.*?)"
        r"(?=(?:Question\s*[:：]\s*|Q\.?\s*[:：]?\s*|$))",
        re.DOTALL | re.IGNORECASE
    ), "qa"),
]

# 匹配编号章节标题
_SECTION_PATTERNS = [
    re.compile(r"^\s*(\d+)\s*[\.．、)\]〕\s]+(.+)$"),       # 1. 标题 / 1、标题
    re.compile(r"^\s*([一二三四五六七八九十]+)[、．.]\s*(.+)$"),  # 一、标题 或 一.标题
    re.compile(r"^\s*[\(（]\s*(\d+)\s*[\)）]\s*(.+)$"),       # (1) 标题
    re.compile(r"^\s*[①②③④⑤⑥⑦⑧⑨⑩]\s*(.+)$"),           # ①标题
    re.compile(r"^\s*#{1,6}\s+(.+)$"),                       # Markdown 标题
]

# 匹配项目符号行
_BULLET_PATTERN = re.compile(r"^\s*[-•*]\s+(.+)$")


def _try_qa_pattern(text: str) -> list[tuple[str, str, str]]:
    """尝试按显式 Q&A 模式拆分"""
    entries = []
    for pattern, _ in _QA_PATTERNS:
        matches = pattern.findall(text)
        if matches:
            for q_raw, s_raw in matches:
                q = q_raw.strip().replace("\n", " ").strip()
                s = s_raw.strip().replace("\n", " ").strip()
                if q and s:
                    entries.append((q, s, ""))
            if entries:
                return entries
    return entries


def _try_section_split(text: str) -> list[tuple[str, str, str]]:
    """按编号章节拆分：标题作为问题，后续内容为答案"""
    lines = text.splitlines()
    entries = []
    current_title = ""
    current_body = []

    for line in lines:
        matched = False
        for pat in _SECTION_PATTERNS:
            m = pat.match(line)
            if m:
                if current_title and current_body:
                    entries.append((current_title.strip(), "\n".join(current_body).strip(), ""))
                groups = m.groups()
                title_text = groups[-1].strip()
                current_title = title_text
                current_body = []
                matched = True
                break
        if not matched:
            stripped = line.strip()
            if stripped:
                current_body.append(stripped)

    if current_title and current_body:
        entries.append((current_title.strip(), "\n".join(current_body).strip(), ""))

    return entries


def _try_bullet_split(text: str) -> list[tuple[str, str, str]]:
    """按项目符号行拆分"""
    lines = text.splitlines()
    entries = []
    current_bullets = []
    current_detail = []

    for line in lines:
        stripped = line.strip()
        if _BULLET_PATTERN.match(line):
            if current_bullets or current_detail:
                entries.append(_finalize_bullet_entry(current_bullets, current_detail))
            current_bullets = []
            current_detail = []
            bullet_text = _BULLET_PATTERN.sub(r"\1", stripped).strip()
            current_bullets.append(bullet_text)
        elif stripped:
            current_detail.append(stripped)

    if current_bullets or current_detail:
        entries.append(_finalize_bullet_entry(current_bullets, current_detail))

    return entries


def _finalize_bullet_entry(bullets: list[str], detail: list[str]) -> tuple[str, str, str]:
    if bullets:
        question = "；".join(bullets[:2])
        solution = "\n".join(bullets + detail)
    else:
        question = detail[0][:60] if detail else ""
        solution = "\n".join(detail)
    return (question, solution, "")


def _try_blank_line_split(text: str) -> list[tuple[str, str, str]]:
    """按空白行拆分为段落块，每两个连续块组成一个问题-答案对"""
    blocks = re.split(r"\n\s*\n", text)
    blocks = [b.strip() for b in blocks if b.strip()]

    if not blocks:
        return []
    if len(blocks) <= 1:
        return [(blocks[0][:80], blocks[0], "")]

    entries = []
    for i in range(0, len(blocks) - 1, 2):
        q = blocks[i][:100]
        s = blocks[i + 1]
        if q and s:
            entries.append((q, s, ""))
    if len(blocks) % 2 == 1:
        remaining = blocks[-1]
        if remaining:
            entries.append((remaining[:80], remaining, ""))

    return entries


def _detect_tags(question: str) -> str:
    """根据问题文本自动推断标签"""
    tags = []
    keyword_map = {
        "安全": ["安全", "密码", "权限", "授权", "认证", "加密", "漏洞", "攻击", "防火墙", "漏洞", "审计"],
        "运维": ["运维", "部署", "重启", "停止", "启动", "监控", "告警", "维护", "备份", "恢复"],
        "网络": ["网络", "DNS", "IP", "端口", "ping", "连接", "断网", "带宽", "路由"],
        "存储": ["磁盘", "存储", "备份", "空间", "清理", "容量", "硬盘", "挂载"],
        "系统": ["系统", "服务器", "Windows", "Linux", "操作系统", "进程", "服务"],
        "数据库": ["数据库", "SQL", "MySQL", "Oracle", "Redis", "MongoDB", "数据"],
        "合规": ["合规", "规范", "制度", "标准", "要求", "流程", "审批", "报告"],
    }
    question_lower = question.lower()
    for tag, keywords in keyword_map.items():
        for kw in keywords:
            if kw.lower() in question_lower:
                tags.append(tag)
                break
    return ",".join(tags[:3])


# ============================================================
# 主入口
# ============================================================

def parse_document_to_entries(file_path: str) -> list[tuple[str, str, str]]:
    """
    从文档文件提取文本并智能拆分为 Q&A 条目。

    图片处理：
    - DOCX：通过 python-docx 提取所有内嵌图片，保存到 KB_IMAGE_DIR
    - PDF：通过 PyMuPDF(fitz) 提取图片（需安装），否则跳过
    - 图片引用以 [IMG: filename] 标记追加到条目方案文本末尾

    拆分策略（按优先级尝试）：
    1. 显式 Q&A 模式（Q:/A:, 问题：/答案：等）
    2. 编号章节拆分（1./2., 一、/二、, # 标题等）
    3. 项目符号拆分（- xxx / * xxx）
    4. 空白行分段配对
    5. 全文作为单一条目

    返回: [(question, solution, tags), ...]
    """
    ext = os.path.splitext(file_path)[1].lower()

    # 1. 提取文本 + 图片
    if ext == ".docx":
        text, images = extract_text_from_docx(file_path)
    elif ext == ".pdf":
        text, images = extract_text_from_pdf(file_path)
    else:
        # 兜底：直接读文本文件
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        images = []

    if not text or not text.strip():
        return []

    # 2. 智能拆分
    entries = _try_qa_pattern(text)
    if not entries:
        entries = _try_section_split(text)
    if not entries:
        entries = _try_bullet_split(text)
    if not entries:
        entries = _try_blank_line_split(text)

    # 3. 追加图片引用到最后一个条目的方案中
    if images:
        image_dir = _get_image_dir()
        image_refs = []
        for idx, (_, filename) in enumerate(images, 1):
            full_path = os.path.join(image_dir, filename)
            image_refs.append(f"  - 图片{idx}: {full_path}")
        image_block = "\n\n---\n[📷 本条文档包含以下图片]\n" + "\n".join(image_refs)

        if entries:
            # 追加到最后一个条目的方案末尾
            q, s, t = entries[-1]
            s = s.rstrip() + image_block
            entries[-1] = (q, s, t)

    # 4. 为每个条目自动推断标签并截断
    enriched = []
    for q, s, _tags in entries:
        if len(q) > 200:
            q = q[:197] + "..."
        if len(s) > 3000:
            s = s[:2997] + "..."
        tag = _detect_tags(q)
        enriched.append((q, s, tag))

    logger.info(
        f"文档 {os.path.basename(file_path)} 解析: {len(entries)} 条条目, "
        f"{len(images)} 张图片"
    )

    return enriched
