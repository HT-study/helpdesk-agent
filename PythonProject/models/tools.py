# models/tools.py - Agent 工具集
# 所有工具统一走 _run_tool：自动写审计日志 + 异常兜底
import datetime
import hashlib
import json
import os
import platform
import shlex
import subprocess
import time
from io import StringIO

import yaml
from langchain.tools import tool

from services.logger_service import get_logger
from config import (
    ALLOWED_WRITE_DIRS,
    DANGEROUS_COMMANDS,
    MAX_FILE_SIZE_BYTES,
    CMD_TIMEOUT_SECONDS,
)

logger = get_logger()
_IS_WIN = platform.system() == "Windows"


# ============================================================
# 统一工具执行入口（审计日志 + 异常兜底）
# ============================================================
def _run_tool(name: str, args: dict, fn):
    """执行工具实现，统一记录审计日志。返回字符串结果。"""
    start = time.time()
    success = True
    result = ""
    try:
        result = fn()
        if not isinstance(result, str):
            result = str(result)
        return result
    except Exception as e:
        success = False
        result = f"❌ 工具 {name} 异常: {e}"
        logger.error(result, exc_info=True)
        return result
    finally:
        try:
            from services.audit_service import audit_log
            audit_log(name, args, result[:2000], success, int((time.time() - start) * 1000))
        except Exception as e:
            logger.warning(f"审计日志写入失败: {e}")


# ============================================================
# 路径安全检查
# ============================================================
def _is_allowed_write_path(file_path: str) -> tuple[bool, str]:
    try:
        norm = os.path.normpath(os.path.abspath(file_path))
    except Exception:
        return False, f"路径无法解析: {file_path}"

    for allowed in ALLOWED_WRITE_DIRS:
        allowed_norm = os.path.normpath(allowed)
        if norm.lower().startswith(allowed_norm.lower()):
            remaining = norm[len(allowed_norm):]
            if remaining == "" or remaining.startswith(os.sep):
                return True, ""
    return False, f"安全限制：仅允许写入 {', '.join(ALLOWED_WRITE_DIRS)}"


def _validate_no_dangerous(cmd: str) -> str | None:
    try:
        tokens = shlex.split(cmd)
    except ValueError:
        tokens = cmd.split()

    for token in tokens:
        stripped = token.strip().lower()
        for dangerous in DANGEROUS_COMMANDS:
            if stripped == dangerous.lower():
                return f"安全拦截：禁止执行包含 '{dangerous}' 的命令"
    return None


def _run_cmd(cmd: str, timeout: int | None = None) -> str:
    """执行系统命令并返回输出（内部通用）"""
    result = subprocess.run(
        cmd, shell=True, capture_output=True, text=True,
        timeout=timeout or CMD_TIMEOUT_SECONDS,
    )
    return result.stdout if result.stdout else result.stderr


# ============================================================
# 基础工具
# ============================================================
@tool
def get_time() -> str:
    """获取当前系统时间"""
    return _run_tool("get_time", {}, lambda: datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))


@tool
def execute_command(cmd: str) -> str:
    """执行系统命令（仅用于学习）"""
    def _impl():
        cmd_preview = cmd if len(cmd) <= 40 else cmd[:20] + "..." + cmd[-20:]
        logger.info(f"调用工具: execute_command, 命令长度: {len(cmd)}")
        logger.debug(f"命令内容: {cmd_preview}")

        block_msg = _validate_no_dangerous(cmd)
        if block_msg:
            logger.warning(block_msg)
            return f"⛔ {block_msg}"

        try:
            output = _run_cmd(cmd)
            logger.info(f"命令执行成功，输出长度: {len(output)}")
            return output
        except subprocess.TimeoutExpired:
            msg = f"命令超时（{CMD_TIMEOUT_SECONDS}s）"
            logger.warning(msg)
            return f"⏱ {msg}"
    return _run_tool("execute_command", {"cmd": cmd}, _impl)


@tool
def read_file(file_path: str) -> str:
    """智能读取文件内容（最大 5MB）"""
    def _impl():
        logger.info(f"调用工具: read_file, 路径: {file_path}")

        if not os.path.exists(file_path):
            return f"❌ 文件不存在: {file_path}"

        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE_BYTES:
            return (
                f"❌ 文件过大: {file_size / 1024 / 1024:.1f} MB，"
                f"超过上限 {MAX_FILE_SIZE_BYTES / 1024 / 1024:.1f} MB"
            )

        ext = os.path.splitext(file_path)[1].lower()
        base_info = f"文件: {os.path.basename(file_path)}\n大小: {file_size} 字节\n"

        if ext in [".txt", ".py", ".log", ".csv", ".md", ".ini", ".conf"]:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="gbk") as f:
                    content = f.read()
            logger.info(f"文本文件读取成功: {len(content)} 字符")
            return base_info + "内容:\n" + content

        elif ext == ".json":
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return base_info + "内容:\n" + json.dumps(data, indent=2, ensure_ascii=False)

        elif ext in [".yaml", ".yml"]:
            with open(file_path, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f)
            return base_info + "内容:\n" + yaml.dump(data, allow_unicode=True)

        elif ext in [".xlsx", ".xls"]:
            import pandas as pd

            excel_file = pd.ExcelFile(file_path)
            result = [f"📊 Excel 文件，包含 {len(excel_file.sheet_names)} 个工作表:"]
            for sheet in excel_file.sheet_names:
                engine = "openpyxl" if ext == ".xlsx" else "xlrd"
                df = pd.read_excel(file_path, sheet_name=sheet, engine=engine)
                result.append(f"\n--- 工作表: {sheet} ---\n{df.to_string()}")
            return base_info + "\n".join(result)

        elif ext == ".docx":
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return base_info + "内容:\n" + "\n".join(paragraphs)

        elif ext == ".pdf":
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return base_info + f"总页数: {len(reader.pages)}\n内容:\n" + text

        elif ext in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"]:
            from PIL import Image

            img = Image.open(file_path)
            return (
                base_info
                + f"图片信息:\n- 尺寸: {img.width} x {img.height}\n- 格式: {img.format}\n- 模式: {img.mode}"
            )

        else:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return (
                base_info
                + f"文件类型: 二进制文件\nMD5: {hash_md5.hexdigest()}\n(内容无法直接显示)"
            )

    return _run_tool("read_file", {"file_path": file_path}, _impl)


@tool
def write_file(file_path: str, content: str) -> str:
    """智能写入文件（根据扩展名自动选择格式）"""
    def _impl():
        logger.info(f"调用工具: write_file, 路径: {file_path}, 内容长度: {len(content)}")

        allowed, msg = _is_allowed_write_path(file_path)
        if not allowed:
            logger.warning(msg)
            return f"⛔ {msg}"

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".json":
            data = json.loads(content)
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            return f"✅ JSON 文件已写入: {file_path}"

        elif ext in [".yaml", ".yml"]:
            data = yaml.safe_load(content)
            with open(file_path, "w", encoding="utf-8") as f:
                yaml.dump(data, f, allow_unicode=True, default_flow_style=False)
            return f"✅ YAML 文件已写入: {file_path}"

        elif ext in [".xlsx", ".xls"]:
            import pandas as pd

            if ext == ".xls":
                return "❌ 写入 .xls 格式较旧，请使用 .xlsx"
            df = pd.read_csv(StringIO(content))
            df.to_excel(file_path, sheet_name="Sheet1", index=False, engine="openpyxl")
            return f"✅ Excel 文件已写入: {file_path}，共 {len(df)} 行"

        elif ext == ".docx":
            from docx import Document

            doc = Document()
            for para in content.split("\n"):
                if para.strip():
                    doc.add_paragraph(para)
            doc.save(file_path)
            return f"✅ Word 文档已写入: {file_path}"

        else:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            return f"✅ 文件已写入: {file_path}"

    return _run_tool("write_file", {"file_path": file_path, "content_len": len(content)}, _impl)


# ============================================================
# 运维专用工具（A1）
# ============================================================
@tool
def list_processes() -> str:
    """列出当前系统进程（按内存占用排序，前 20 个）"""
    def _impl():
        try:
            import psutil
            procs = []
            for p in psutil.process_iter(["pid", "name", "memory_percent", "cpu_percent", "username"]):
                try:
                    procs.append(p.info)
                except Exception:
                    pass
            procs.sort(key=lambda x: x.get("memory_percent") or 0, reverse=True)
            lines = ["PID\t名称\t\t内存%\tCPU%\t用户"]
            for p in procs[:20]:
                lines.append(
                    f"{p['pid']}\t{p['name']}\t{p.get('memory_percent') or 0:.1f}\t"
                    f"{p.get('cpu_percent') or 0:.1f}\t{p.get('username') or '-'}"
                )
            return "\n".join(lines)
        except ImportError:
            cmd = "tasklist /fo csv" if _IS_WIN else "ps aux --sort=-%mem | head -n 20"
            out = _run_cmd(cmd)
            return out[:5000]

    return _run_tool("list_processes", {}, _impl)


@tool
def kill_process(pid: int) -> str:
    """结束指定 PID 的进程（危险操作；系统关键进程会被拒绝）"""
    def _impl():
        # 保护系统关键进程
        protected = {0, 1, 4}
        if pid in protected:
            return f"⛔ 拒绝结束关键系统进程 PID={pid}"
        cmd = f"taskkill /PID {pid} /F" if _IS_WIN else f"kill -9 {pid}"
        try:
            out = _run_cmd(cmd, timeout=10)
            return f"✅ 已请求结束进程 {pid}\n{out[:500]}"
        except Exception as e:
            return f"❌ 结束进程失败: {e}"

    return _run_tool("kill_process", {"pid": pid}, _impl)


@tool
def list_services() -> str:
    """列出系统服务及其状态"""
    def _impl():
        if _IS_WIN:
            cmd = 'sc query state= all | findstr /R "SERVICE_NAME STATE"'
        else:
            cmd = "systemctl list-units --type=service --no-pager"
        out = _run_cmd(cmd)
        return out[:8000]

    return _run_tool("list_services", {}, _impl)


@tool
def network_check(host: str) -> str:
    """网络诊断：DNS 解析 + ping 连通性"""
    def _impl():
        import socket
        lines = []
        # DNS
        try:
            ip = socket.gethostbyname(host)
            lines.append(f"DNS: {host} → {ip}")
        except Exception as e:
            return f"❌ DNS 解析失败: {e}"
        # ping
        cmd = f"ping -n 4 {host}" if _IS_WIN else f"ping -c 4 {host}"
        try:
            out = _run_cmd(cmd, timeout=20)
            lines.append("Ping 结果:\n" + out[-1500:])
        except Exception as e:
            lines.append(f"Ping 失败: {e}")
        return "\n".join(lines)

    return _run_tool("network_check", {"host": host}, _impl)


@tool
def system_stats() -> str:
    """获取系统资源快照：CPU、内存、磁盘"""
    def _impl():
        try:
            import psutil
            cpu = psutil.cpu_percent(interval=1)
            mem = psutil.virtual_memory()
            disk = psutil.disk_usage("/")
            return (
                f"CPU 使用率: {cpu}%\n"
                f"内存: {mem.percent}% (已用 {mem.used // 1024 // 1024} MB / 共 {mem.total // 1024 // 1024} MB)\n"
                f"磁盘: {disk.percent}% (已用 {disk.used // 1024 // 1024 // 1024} GB / 共 {disk.total // 1024 // 1024 // 1024} GB)"
            )
        except ImportError:
            cmd = "wmic CPU get LoadPercentage /value" if _IS_WIN else "free -m && df -h"
            return _run_cmd(cmd)[:3000]

    return _run_tool("system_stats", {}, _impl)


@tool
def search_log(file_path: str, keyword: str, max_lines: int = 50) -> str:
    """在日志文件中检索关键字，返回匹配行（最多 max_lines 条）"""
    def _impl():
        if not os.path.exists(file_path):
            return f"❌ 文件不存在: {file_path}"
        size = os.path.getsize(file_path)
        if size > MAX_FILE_SIZE_BYTES * 4:
            return f"❌ 文件过大（{size // 1024 // 1024} MB），超过检索上限"

        matches = []
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    if keyword in line:
                        matches.append(line.rstrip())
                        if len(matches) >= max_lines:
                            break
        except Exception as e:
            return f"❌ 读取失败: {e}"

        if not matches:
            return f"未找到包含 '{keyword}' 的行"
        return f"找到 {len(matches)} 条匹配:\n" + "\n".join(matches)

    return _run_tool("search_log", {"file_path": file_path, "keyword": keyword, "max_lines": max_lines}, _impl)


@tool
def list_dir(path: str) -> str:
    """列出目录内容（含类型、大小、修改时间，最多 200 项）"""
    def _impl():
        if not os.path.isdir(path):
            return f"❌ 不是目录: {path}"
        items = []
        for name in sorted(os.listdir(path))[:200]:
            full = os.path.join(path, name)
            try:
                st = os.stat(full)
                kind = "D" if os.path.isdir(full) else "F"
                size = st.st_size if not os.path.isdir(full) else 0
                mtime = datetime.datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M")
                items.append(f"[{kind}] {size:>10}  {mtime}  {name}")
            except Exception:
                items.append(f"[?]         -         {name}")
        return f"目录: {path} （{len(items)} 项，已截断）\n" + "\n".join(items)

    return _run_tool("list_dir", {"path": path}, _impl)


@tool
def search_kb(query: str) -> str:
    """检索知识库中过往的问题解决方案（语义搜索 + 关键词兜底）"""
    def _impl():
        from services.kb_service import search_semantic, search_kb as keyword_search
        results = search_semantic(query, limit=5)
        if not results:
            results = keyword_search(query, limit=5)
        if not results:
            return f"知识库中未找到与 '{query}' 相关的条目"
        out = [f"📚 找到 {len(results)} 条相关知识（语义匹配）:"]
        for i, r in enumerate(results, 1):
            out.append(f"\n--- {i}. 问题: {r['question']}")
            out.append(f"   方案: {r['solution']}")
            if r.get("tags"):
                out.append(f"   标签: {r['tags']}")
            out.append(f"   匹配度: {r['score']:.1%}")
        return "\n".join(out)

    return _run_tool("search_kb", {"query": query}, _impl)


# ============================================================
# 远程主机工具（B2 - 多主机管理）
# ============================================================
@tool
def list_hosts() -> str:
    """列出已配置的远程主机"""
    def _impl():
        from services.host_service import list_hosts as _list
        hosts = _list()
        if not hosts:
            return "未配置远程主机，请先到 /hosts 页面添加"
        out = ["📋 已配置远程主机:"]
        for h in hosts:
            out.append(f"  - {h['label']} ({h['hostname']}:{h['port']}, {h['username']}, {h['auth_type']})")
        out.append("\n使用 execute_on_host(host_label, cmd) 在指定主机执行命令")
        return "\n".join(out)
    return _run_tool("list_hosts", {}, _impl)


@tool
def execute_on_host(host_label: str, cmd: str) -> str:
    """在指定远程主机上执行命令（参数：host_label为主机标签，cmd为命令）"""
    def _impl():
        from services.host_service import execute_on_host as _exec
        success, output = _exec(host_label, cmd)
        if success:
            return f"✅ [{host_label}] 执行结果:\n{output}"
        return f"❌ [{host_label}] 执行失败: {output}"
    return _run_tool("execute_on_host", {"host_label": host_label, "cmd": cmd}, _impl)


@tool
def run_on_all_hosts(cmd: str) -> str:
    """在所有已配置的远程主机上并行执行命令"""
    def _impl():
        from services.host_service import run_on_all_hosts as _run_all
        results = _run_all(cmd)
        out = []
        for label, success, output in results:
            if not label:
                return output
            status = "✅" if success else "❌"
            body = output[:2000]
            out.append(f"{status} [{label}]:\n{body}\n")
        return "\n".join(out)
    return _run_tool("run_on_all_hosts", {"cmd": cmd}, _impl)


@tool
def check_host_health(host_label: str) -> str:
    """检查远程主机连通性（TCP端口 + SSH登录）"""
    def _impl():
        from services.host_service import check_host_health as _check
        ok, msg = _check(host_label)
        return f"{'✅' if ok else '❌'} [{host_label}] {msg}"
    return _run_tool("check_host_health", {"host_label": host_label}, _impl)


# ============================================================
# Excel 数据分析工具
# ============================================================
_EXCEL_CHART_SEQ = 0


def _chart_path():
    """生成唯一的图表文件路径"""
    import time
    global _EXCEL_CHART_SEQ
    _EXCEL_CHART_SEQ += 1
    from config import CHART_DIR
    os.makedirs(CHART_DIR, exist_ok=True)
    return os.path.join(CHART_DIR, f"chart_{int(time.time())}_{_EXCEL_CHART_SEQ}.png")


def _read_excel(file_path, sheet=None):
    """读取 Excel 并返回 (df, sheet_name, 错误信息)"""
    import pandas as pd
    if not os.path.exists(file_path):
        return None, None, f"❌ 文件不存在: {file_path}"
    if file_size := os.path.getsize(file_path) > MAX_FILE_SIZE_BYTES * 4:
        return None, None, f"❌ 文件过大（{file_size // 1024 // 1024} MB）"
    try:
        xls = pd.ExcelFile(file_path)
        if sheet is None:
            sheet = xls.sheet_names[0]
        df = xls.parse(sheet)
        return df, sheet, None
    except Exception as e:
        return None, None, f"❌ 读取失败: {e}"


def _truncate_df(df, max_rows=10, max_cols=15):
    """截断 DataFrame 为可显示大小"""
    if df.shape[1] > max_cols:
        df = df.iloc[:, :max_cols]
    show = df.head(max_rows)
    truncated = len(df) > max_rows
    return show, truncated


@tool
def excel_summary(file_path: str, sheet: str = "") -> str:
    """分析 Excel 文件结构：工作表列表、列信息、数值统计、空值率"""
    def _impl():
        import pandas as pd
        if not os.path.exists(file_path):
            return f"❌ 文件不存在: {file_path}"
        try:
            xls = pd.ExcelFile(file_path)
        except Exception as e:
            return f"❌ 读取失败: {e}"

        lines = [f"📊 文件: {os.path.basename(file_path)}"]
        sheets = xls.sheet_names
        lines.append(f"工作表: {sheets} ({len(sheets)} 个)")

        targets = [sheet] if sheet else sheets
        for s_name in targets:
            try:
                df = xls.parse(s_name)
                lines.append(f"\n--- 工作表: {s_name} ({df.shape[0]} 行 × {df.shape[1]} 列) ---")
                if df.shape[0] == 0:
                    lines.append("（空表）")
                    continue
                # 列信息
                header = f"{'列名':<20} {'类型':<12} {'空值率':<10} {'唯一值':<8}"
                lines.append(header)
                lines.append("-" * 50)
                for col in df.columns[:20]:
                    col_str = str(col)[:20]
                    dtype = str(df[col].dtype)[:10]
                    null_pct = f"{df[col].isnull().mean() * 100:.0f}%"
                    nunique = df[col].nunique()
                    lines.append(f"{col_str:<20} {dtype:<12} {null_pct:<10} {nunique:<8}")
                # 数值列统计
                num_cols = df.select_dtypes(include='number').columns[:5]
                if len(num_cols) > 0:
                    lines.append(f"\n数值列统计（前 {len(num_cols)} 列）:")
                    lines.append(f"{'列名':<20} {'均值':>12} {'最小':>12} {'最大':>12} {'标准差':>12}")
                    lines.append("-" * 68)
                    for col in num_cols:
                        c = df[col].dropna()
                        if len(c) > 0:
                            lines.append(
                                f"{str(col):<20} {c.mean():>12.2f} {c.min():>12.2f} {c.max():>12.2f} {c.std():>12.2f}"
                            )
            except Exception as e:
                lines.append(f"工作表 {s_name} 分析失败: {e}")
        return "\n".join(lines)

    return _run_tool("excel_summary", {"file_path": file_path, "sheet": sheet}, _impl)


@tool
def excel_filter(file_path: str, column: str, operator: str, value: str, limit: int = 10, sheet: str = "") -> str:
    """按条件筛选 Excel 行，支持 operator: ==, !=, >, <, >=, <=, contains, in"""
    def _impl():
        import pandas as pd
        df, s_name, err = _read_excel(file_path, sheet or None)
        if err:
            return err
        if column not in df.columns:
            return f"❌ 列 '{column}' 不存在，可用列: {', '.join(str(c) for c in df.columns[:20])}"

        try:
            if operator == "contains":
                mask = df[column].astype(str).str.contains(value, case=False, na=False)
            elif operator == "in":
                vals = [v.strip() for v in value.split(",")]
                mask = df[column].astype(str).str.lower().isin([v.lower() for v in vals])
            else:
                col_type = df[column].dtype
                if pd.api.types.is_numeric_dtype(col_type):
                    v = float(value)
                elif pd.api.types.is_datetime64_any_dtype(col_type):
                    v = pd.Timestamp(value)
                else:
                    v = value
                ops = {">": ">", "<": "<", ">=": ">=", "<=": "<=", "==": "==", "!=": "!="}
                if operator not in ops:
                    return f"❌ 不支持的运算符: {operator}，支持: {', '.join(ops.keys())}"
                mask = eval(f"df[col_type_v] {ops[operator]} df[col_type_v]", {"df": df, "col_type_v": column, "v": v})
                # Actually build the query properly
                if operator == "==":
                    mask = df[column] == v
                elif operator == "!=":
                    mask = df[column] != v
                elif operator == ">":
                    mask = df[column] > v
                elif operator == "<":
                    mask = df[column] < v
                elif operator == ">=":
                    mask = df[column] >= v
                elif operator == "<=":
                    mask = df[column] <= v

            result = df[mask]
            if len(result) == 0:
                return f"未找到满足条件 {column} {operator} {value} 的行"
            show, truncated = _truncate_df(result, limit)
            out = [f"筛选: {column} {operator} {value} → 匹配 {len(result)} 行" + (f"，显示前 {limit} 行" if truncated else "")]
            out.append(show.to_string(index=False))
            return "\n".join(out)
        except Exception as e:
            return f"❌ 筛选失败: {e}"

    return _run_tool("excel_filter", {"file_path": file_path, "column": column, "operator": operator, "value": value, "limit": limit, "sheet": sheet}, _impl)


@tool
def excel_aggregate(file_path: str, group_col: str, agg_col: str, agg_func: str = "sum", sheet: str = "") -> str:
    """按分组列聚合统计，agg_func 支持: sum, mean, count, max, min, median, std"""
    def _impl():
        import pandas as pd
        df, s_name, err = _read_excel(file_path, sheet or None)
        if err:
            return err
        if group_col not in df.columns:
            return f"❌ 分组列 '{group_col}' 不存在"
        if agg_col and agg_col not in df.columns:
            return f"❌ 聚合列 '{agg_col}' 不存在"
        funcs = {"sum": "sum", "mean": "mean", "count": "count", "max": "max", "min": "min", "median": "median", "std": "std"}
        if agg_func not in funcs:
            return f"❌ 不支持的聚合函数: {agg_func}，支持: {', '.join(funcs.keys())}"

        try:
            if agg_col:
                result = df.groupby(group_col)[agg_col].agg(agg_func).reset_index()
            else:
                num_cols = df.select_dtypes(include='number').columns[:5]
                result = df.groupby(group_col)[num_cols].agg(agg_func).reset_index()

            result = result.sort_values(result.columns[-1], ascending=False)
            show, truncated = _truncate_df(result, 20)
            out = [f"按 {group_col} 分组, 对 {agg_col or '数值列'} 求 {agg_func}:"]
            out.append(show.to_string(index=False))
            if truncated:
                out.append(f"... 共 {len(result)} 行")
            return "\n".join(out)
        except Exception as e:
            return f"❌ 聚合失败: {e}"

    return _run_tool("excel_aggregate", {"file_path": file_path, "group_col": group_col, "agg_col": agg_col, "agg_func": agg_func, "sheet": sheet}, _impl)


@tool
def excel_chart(file_path: str, x_col: str, y_col: str, chart_type: str = "bar", sheet: str = "") -> str:
    """生成 Excel 数据图表（bar/line/pie），返回图表图片 URL"""
    def _impl():
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import pandas as pd

        # 中文字体配置
        for font_name in ["SimHei", "Microsoft YaHei", "PingFang SC", "Noto Sans CJK"]:
            try:
                matplotlib.font_manager.findfont(font_name, fallback_to_default=False)
                plt.rcParams['font.sans-serif'] = [font_name]
                break
            except Exception:
                continue
        plt.rcParams['axes.unicode_minus'] = False

        df, s_name, err = _read_excel(file_path, sheet or None)
        if err:
            return err
        if x_col not in df.columns:
            return f"❌ X 列 '{x_col}' 不存在"
        if y_col and y_col not in df.columns:
            return f"❌ Y 列 '{y_col}' 不存在"

        try:
            if chart_type == "pie":
                # 饼图只取前 10 项
                data = df[[x_col, y_col]].dropna().head(10)
                labels = data[x_col].astype(str).tolist()
                values = pd.to_numeric(data[y_col], errors='coerce').fillna(0).tolist()
                fig, ax = plt.subplots(figsize=(8, 6))
                ax.pie(values, labels=labels, autopct='%1.1f%%')
                ax.set_title(f"{os.path.basename(file_path)} - {x_col} 分布")
            elif chart_type == "line":
                data = df[[x_col, y_col]].dropna().sort_values(by=x_col)
                fig, ax = plt.subplots(figsize=(10, 5))
                ax.plot(data[x_col].astype(str), pd.to_numeric(data[y_col], errors='coerce'), marker='o')
                ax.set_xlabel(x_col)
                ax.set_ylabel(y_col)
                ax.tick_params(axis='x', rotation=45)
                ax.set_title(f"{os.path.basename(file_path)} - {x_col} vs {y_col}")
            else:  # bar
                data = df.groupby(x_col)[y_col].sum().reset_index().head(20)
                fig, ax = plt.subplots(figsize=(10, 5))
                ax.bar(data[x_col].astype(str), pd.to_numeric(data[y_col], errors='coerce'))
                ax.set_xlabel(x_col)
                ax.set_ylabel(y_col)
                ax.tick_params(axis='x', rotation=45)
                ax.set_title(f"{os.path.basename(file_path)} - {x_col} 按 {y_col} 汇总")

            plt.tight_layout()
            path = _chart_path()
            plt.savefig(path, dpi=120)
            plt.close(fig)
            return f"✅ 图表已生成: /file/chart/{os.path.basename(path)}"

        except Exception as e:
            return f"❌ 图表生成失败: {e}"

    return _run_tool("excel_chart", {"file_path": file_path, "x_col": x_col, "y_col": y_col, "chart_type": chart_type, "sheet": sheet}, _impl)